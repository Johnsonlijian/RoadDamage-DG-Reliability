from __future__ import annotations

import importlib.util
import re
import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch


ROOT = Path(__file__).resolve().parents[1]
TODAY = "2026-05-22"
SRC_PKG = ROOT / "submission_package" / "AEI_RoadDamageDG_2026-05-17"
PKG = ROOT / "submission_package" / "JCICE_RoadDamageDG_2026-05-22"
ROUND = ROOT / f"rounds/R35_jcice_technical_reframe_{TODAY}"
ROUND.mkdir(parents=True, exist_ok=True)
FIG_DIR = PKG / "figures_enhanced"
SRC_TABLES = PKG / "source_tables"
V21 = SRC_PKG / "manuscript_v21_r34_rtdetr_validation.md"
V22 = PKG / "manuscript_v22_jcice_technical_paper.md"
V22_DOCX = PKG / "Manuscript_JCICE_v22_technical_paper.docx"
REPO_URL = "https://github.com/Johnsonlijian/RoadDamage-DG-Reliability"


def load_v19_builder():
    path = ROOT / "scripts" / "58_build_v19_methodology_guidance.py"
    spec = importlib.util.spec_from_file_location("v19_builder", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Cannot load {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    module.PKG = PKG
    return module


v19 = load_v19_builder()


def prepare_package() -> None:
    PKG.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    SRC_TABLES.mkdir(parents=True, exist_ok=True)
    for src in (SRC_PKG / "figures_enhanced").glob("*"):
        if src.is_file():
            shutil.copy2(src, FIG_DIR / src.name)
    for src in (SRC_PKG / "source_tables").glob("*"):
        if src.is_file():
            shutil.copy2(src, SRC_TABLES / src.name)


def make_framework_figure() -> Path:
    path = FIG_DIR / "fig01_jcice_domain_aware_reliability_audit_framework.png"
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 8.2,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.spines.left": False,
            "axes.spines.bottom": False,
        }
    )
    fig, ax = plt.subplots(figsize=(12.4, 5.7), constrained_layout=True)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")

    columns = [
        {
            "title": "Data and task boundary",
            "x": 0.35,
            "color": "#DCEBFA",
            "items": [
                "RDD2022 public archive",
                "47,420 images",
                "Seven country/capture domains",
                "38,385 XML annotations",
                "D00/D10/D20/D40 supervised task",
                "10,705 additional-label boxes retained",
            ],
        },
        {
            "title": "Reliability-audit pipeline",
            "x": 4.05,
            "color": "#E6F3E8",
            "items": [
                "1. Inventory domains and labels",
                "2. Ordinary mixed-domain validation",
                "3. Leave-one-domain-out validation",
                "4. Source-domain budget sweep",
                "5. Detector-family checks",
                "6. Prediction export and calibration",
                "7. Confidence-frontier analysis",
            ],
        },
        {
            "title": "Reliability-boundary outputs",
            "x": 7.95,
            "color": "#F7E8D8",
            "items": [
                "Ordinary-to-LODO gap",
                "Weak held-out domains",
                "Weak classes and domain cells",
                "Detector-family sensitivity",
                "High-confidence calibration gap",
                "Precision-coverage frontier",
                "Six-boundary reporting standard",
            ],
        },
    ]

    for col in columns:
        box = FancyBboxPatch(
            (col["x"], 1.05),
            3.35,
            5.35,
            boxstyle="round,pad=0.05,rounding_size=0.08",
            linewidth=1.2,
            facecolor=col["color"],
            edgecolor="#333333",
        )
        ax.add_patch(box)
        ax.text(col["x"] + 1.675, 6.05, col["title"], ha="center", va="center", fontsize=10.5, weight="bold")
        y = 5.55
        for item in col["items"]:
            ax.text(col["x"] + 0.22, y, u"\u2022 " + item, ha="left", va="top", fontsize=8.25)
            y -= 0.58

    for x1, x2 in [(3.72, 4.02), (7.45, 7.92)]:
        ax.add_patch(
            FancyArrowPatch(
                (x1, 3.75),
                (x2, 3.75),
                arrowstyle="-|>",
                mutation_scale=16,
                linewidth=1.4,
                color="#333333",
            )
        )

    footer = (
        "Detector scores are treated as civil-infrastructure screening evidence only after "
        "domain, label, budget, class, detector-family, calibration, and confidence-frontier boundaries are visible."
    )
    ax.text(6.0, 0.45, footer, ha="center", va="center", fontsize=9.2, color="#222222")
    fig.suptitle("Domain-aware reliability-audit framework for infrastructure image detection", fontsize=13, weight="bold")
    fig.savefig(path, dpi=300)
    fig.savefig(path.with_suffix(".svg"))
    plt.close(fig)
    return path


def replace_between(text: str, start: str, end: str, replacement: str) -> str:
    pattern = re.escape(start) + r"\n\n.*?\n\n" + re.escape(end)
    return re.sub(pattern, replacement.rstrip() + "\n\n" + end, text, flags=re.S)


def abstract() -> str:
    return (
        "Infrastructure image detectors are increasingly evaluated with public datasets, but ordinary mixed-domain "
        "validation can misrepresent whether a detector score is reliable for a new country, capture platform, or "
        "inspection context. This Technical Paper develops a domain-aware reliability-audit framework for civil-"
        "infrastructure image detection and demonstrates it with RDD2022 road-damage detection as a reproducible "
        "validation-boundary study. The audit first separates the data and task boundary, indexing 47,420 images, "
        "38,385 Pascal VOC XML annotations, 55,007 supervised D00/D10/D20/D40 boxes, and 10,705 additional XML boxes "
        "that define the label boundary. It then compares ordinary mixed-domain validation with leave-one-domain-out "
        "(LODO) validation, tests source-domain budget sensitivity, checks detector-family specificity with YOLOv8s, "
        "Faster R-CNN, RetinaNet, and RT-DETR-L, and audits calibration and confidence-frontier behavior. Under a "
        "five-seed YOLOv8n subset baseline, ordinary validation reaches 0.0853 mAP50 (95% CI 0.0676-0.1031), whereas "
        "mean LODO validation reaches 0.0619 mAP50 (95% CI 0.0553-0.0685). In the YOLOv8s budget sweep, the ordinary-"
        "to-LODO gap rises from 0.0281 to 0.1093 mAP50 as source-domain sampling increases from 80 to 640 images per "
        "source domain. RT-DETR-L improves the subset endpoint to 0.3489 ordinary and 0.2635 mean LODO mAP50, but "
        "India remains weak at 0.0668 mAP50. Pooled LODO predictions do not reach a 0.10 precision floor in the "
        "tested confidence-threshold grid. These results do not establish a deployable road-damage detector; they "
        "establish a reproducible reliability-audit procedure showing how detector scores should be bounded before "
        "being interpreted as civil-infrastructure screening evidence."
    )


def introduction() -> str:
    return """## 1. Introduction

Image-based detection is becoming a routine computing component in civil-infrastructure inspection. Road agencies and researchers can collect pavement images using smartphones, vehicle-mounted cameras, drones, and other mobile platforms, and public datasets such as RDD2022 and CRDDC-2022 make road-damage detection reproducible across countries and acquisition conditions [1-3]. Prior civil-infrastructure vision and pavement-distress studies show that image-based recognition is technically useful, but also sensitive to acquisition setting, distress definition, and validation design [4-7]. For civil-engineering decision support, therefore, a detector score is not enough by itself. The score must be interpretable as screening evidence for inspection prioritization, maintenance planning, or local validation.

The central computing problem is validation-target mismatch. Ordinary mixed-domain validation tests whether a detector recognizes damage when training and validation samples share countries or capture conditions. New-region or new-platform screening asks a different question: whether evidence remains reliable when a full domain is withheld. In civil infrastructure, this distinction matters because a misleading validation number can affect maintenance triage and resource allocation rather than only benchmark ranking. The broader domain-generalization and domain-adaptive object-detection literature shows that benchmark design and target-domain shifts can change conclusions [8-11], but civil-infrastructure image detection still needs a concrete reporting procedure for domain, label, budget, class, detector-family, calibration, and threshold boundaries.

RDD2022 is a useful public corpus for such an audit because its structure exposes several validation boundaries. The extracted archive can be organized into seven country/capture domains, including China_Drone and China_MotorBike as separate acquisition conditions. Its XML annotations include four supervised distress labels, D00, D10, D20, and D40, as well as additional observed XML labels. The file-list structure also separates XML-annotated folders from image-only test folders. These properties make RDD2022 more than a road-damage dataset; it is a controlled setting for studying how task definitions and domain withholding change the engineering meaning of detector scores.

This paper proposes a domain-aware reliability audit rather than a new detector. The audit has six linked modules: domain and label-boundary inventory, ordinary mixed-domain validation, leave-one-domain-out (LODO) validation, source-domain budget sensitivity, detector-family checks, and prediction-export diagnostics covering class/domain errors, calibration, and confidence frontiers. The framework treats road-damage detection as a case study for infrastructure image detection more generally.

The study makes four contributions:

1. A domain-aware reliability-audit framework for infrastructure image detection under country/capture-domain shift.
2. A reproducible RDD2022 validation-boundary study separating mixed-domain recognition from domain-withheld evidence.
3. A multi-layer evidence audit combining source-domain budget, detector-family, class/domain, calibration, and confidence-frontier diagnostics.
4. A six-boundary reporting standard for interpreting detector scores as civil-infrastructure screening evidence."""


def framework_section() -> str:
    return """## 3. Domain-Aware Reliability Audit Framework

### 3.1 Framework overview

The proposed framework treats an infrastructure image-detection study as an evidence-audit problem. The inputs are an image dataset, domain identifiers, a label schema, detector families, and prediction exports. The audit modules are: domain inventory and label-boundary audit; ordinary mixed-domain validation; LODO validation; source-domain budget sweep; detector-family sensitivity check; class and domain diagnostics; and calibration and confidence-frontier analysis. The outputs are an ordinary-vs-LODO reliability gap, weak-domain and weak-class evidence, detector-family sensitivity evidence, a high-confidence calibration gap, a precision-coverage frontier, and a six-boundary reporting table.

![Figure 1. Domain-aware reliability-audit framework for infrastructure image detection.](figures_enhanced/fig01_jcice_domain_aware_reliability_audit_framework.png)

### 3.2 Mixed-domain and domain-withheld validation

Two validation targets are separated. The ordinary setting samples training and validation images from the available annotated domains without withholding a full country/capture domain. It is useful as a mixed-domain recognition reference but is not a new-domain reliability test. The LODO setting withholds one full domain at a time. For each of the seven domains, the detector is trained on the other six domains and evaluated on the held-out domain. This setting tests country/capture transfer within the frozen subset design.

### 3.3 Source-domain budget sweep

The primary repeated baseline uses YOLOv8n pretrained weights [21], image size 320, four epochs, batch size 8, and a frozen subset policy. The same setting is repeated over five random seeds. A second YOLOv8s budget audit uses image size 640 and eight epochs to test whether additional source-domain sampling changes the validation-boundary pattern. The sampled source-domain budgets are 80, 160, 320, and 640 images per source domain.

### 3.4 Detector-family checks

The detector-family layer reduces the risk that the observed validation boundary is only a YOLOv8-family artifact. Faster R-CNN MobileNetV3-320-FPN uses COCO pretrained weights and a four-class prediction head [16]. Each ordinary or LODO setting is trained for eight epochs with batch size 2 and learning rate 0.0025. RetinaNet ResNet50-FPN uses COCO pretrained weights and a four-class dense classification head [17]. Each setting is trained for four epochs with batch size 1 and learning rate 0.001. RT-DETR-L is used as a transformer-family check in the DETR lineage, using the Ultralytics implementation and COCO pretrained weights [18,19]. Each ordinary or LODO setting is trained for eight epochs at 640 px with batch size 2. These checks are not tuned detector leaderboards; they test whether the ordinary-vs-LODO boundary remains visible when the detector family changes.

### 3.5 Prediction-export, calibration, and confidence-frontier diagnostics

Aggregate mAP does not show whether confidence filtering is usable for a civil-infrastructure review queue. Prediction-level exports therefore record true positives, false positives, false negatives, confidence, class, IoU, and domain. The decision frontier reports, for each confidence threshold, the precision among accepted predictions and the coverage retained among prediction rows. Calibration diagnostics report full-stream ECE and high-confidence calibration gaps, following the broader calibration, uncertainty-under-shift, out-of-distribution detection, and selective-prediction literature [12-15]. These diagnostics localize evidence boundaries; they are not operational road-agency referral policies.

### 3.6 Six-boundary reporting standard

A detector score is not treated as civil-infrastructure screening evidence unless six reporting boundaries are visible.

| boundary | minimum report item | engineering use |
| --- | --- | --- |
| Domain | Country/capture platform split; validation must identify held-out domains | Prevents mixed-domain scores from being read as new-region evidence |
| Label | Supervised labels, extra labels, and any merge policy | Keeps damage definitions comparable across studies |
| Budget | Source-domain images/boxes, epochs, image size, and seeds | Separates data scaling from cross-domain reliability |
| Class | Per-class AP or TP/FP/FN diagnostics with support | Identifies which distress categories cannot support decisions |
| Detector family | At least one family-diversity check or a declared architecture limit | Reduces architecture-specific overinterpretation |
| Calibration and confidence frontier | Calibration diagnostic, threshold grid, and precision-coverage table | Shows whether confidence scores can support review queues |"""


def discussion() -> str:
    return """## 5. Discussion

### 5.1 Validation-target mismatch as a civil-infrastructure computing problem

The main implication is methodological. Ordinary mixed-domain validation and LODO validation answer different computing questions. Ordinary validation measures recognition under shared dataset context, whereas LODO validation measures whether evidence transfers when a country/capture domain is withheld. In civil infrastructure, these targets should not be conflated because detector outputs may be used to prioritize inspection, allocate maintenance resources, or justify local validation requirements. The result is consistent with broader domain-generalization findings from WILDS and DomainBed [8,9], but the present study translates the issue into an object-detection and infrastructure-screening setting.

### 5.2 Detector confidence is not engineering confidence

The confidence-frontier and calibration results show why detector confidence should not be read as engineering confidence. In the canonical audit, ordinary validation reaches a 0.30 precision floor only at 0.0026 prediction coverage, while pooled LODO never reaches a 0.10 precision floor in the tested grid. The high-confidence calibration gap also grows under domain withholding. A confidence threshold can therefore be reported as a diagnostic frontier, but it should not be converted into a road-agency referral policy without local validation, calibration, and workload modeling.

### 5.3 Reporting boundaries for accountable infrastructure AI

For accountable infrastructure AI, the practical contribution is a reporting discipline. Domain, label, budget, class, detector-family, calibration, and confidence-frontier boundaries are not formatting details; they determine whether a detector score can be interpreted as screening evidence. The R34 RT-DETR-L validation strengthens this point because the transformer-family check improves absolute ordinary and mean-LODO mAP50 while still leaving India below 0.10 mAP50. The issue is therefore not simply which detector is strongest. The issue is where the validation evidence stops supporting an infrastructure decision."""


def limitations() -> str:
    return """## 6. Limitations

The detector-family checks are used to reduce architecture-specificity risk, not to establish a tuned detector leaderboard. The evidence is subset-scale and should not be compared with full-scale challenge leaderboards or used to rank detector architectures. Faster R-CNN and RT-DETR-L are trained for eight epochs, and RetinaNet is trained for four epochs; none is a tuned multi-seed detector-family benchmark.

Domain adaptation and test-time adaptation are not benchmarked because the purpose is to audit the reliability boundary of validation evidence before adaptation is introduced. Future work should evaluate whether adaptation methods shrink the ordinary-to-LODO gap without obscuring the same reporting boundaries.

RDD2022 is used because its country/capture-domain structure and label annotations allow a controlled public-corpus audit. Extension to additional infrastructure image datasets is needed before claiming task-universal behavior.

The tested detectors should not be interpreted as deployment-ready systems under the reported LODO results. The supervised task is restricted to D00, D10, D20, and D40. Additional XML labels are retained as label-boundary evidence but are not modeled. The available file-list structure provides XML annotations for training folders but not for the image-only test folders, so the experiments use annotated-domain validation rather than official hidden-test challenge evaluation.

Calibration and decision-frontier diagnostics are based on canonical prediction exports and should be treated as evidence-localization diagnostics rather than final operating policies. They do not model agency workload, inspection cost, local acceptance thresholds, or safety consequences. Domain-descriptor correlations use only seven domains and therefore serve as screening cues rather than causal evidence."""


def conclusion_and_practical() -> str:
    return """## 7. Conclusions

This study develops a reproducible domain-aware reliability-audit framework for civil-infrastructure image detection and demonstrates it through RDD2022 road-damage detection. The case study shows that ordinary mixed-domain validation, LODO validation, source-domain scaling, detector-family checking, class/domain diagnostics, calibration, and confidence-frontier analysis produce different evidence boundaries. RT-DETR-L improves the absolute subset endpoint, but the domain-withheld reliability boundary remains visible, with India still weak at 0.0668 mAP50.

The main contribution is a validation protocol and reporting standard rather than a new detector. Detector scores become technically and operationally ambiguous when domain, label, budget, class, detector-family, calibration, and confidence-frontier boundaries are omitted. For computing in civil engineering, the implication is that validation protocols should report not only whether a detector recognizes damage, but also where its evidence stops being interpretable as infrastructure screening support.

## Practical Applications

The framework can be used by civil-infrastructure researchers, road agencies, and technology vendors when evaluating image-based damage detection systems. A mixed-domain mAP value should not be treated as evidence for new-region deployment unless a domain-withheld validation result is also reported. Procurement or pilot studies should request domain definitions, label-boundary rules, source-data budget, per-class diagnostics, detector-family sensitivity, calibration diagnostics, and confidence-frontier tables. For the present RDD2022 case, the tested detectors are not deployment-ready under LODO evaluation. The practical use of the results is therefore to define a reproducible audit checklist for deciding what additional local validation, calibration, or adaptation evidence is needed before detector outputs are used for inspection triage."""


def cover_letter() -> tuple[Path, Path]:
    md = PKG / "CoverLetter_JCICE_v22_technical_paper.md"
    docx = PKG / "CoverLetter_JCICE_v22_technical_paper.docx"
    text = f"""# Cover Letter: Journal of Computing in Civil Engineering

Dear Editor,

We submit the manuscript titled "Domain-Aware Reliability Auditing for Infrastructure Image Detection: A Validation-Boundary Study of RDD2022 Road-Damage Detection" as a Technical Paper for the Journal of Computing in Civil Engineering.

The manuscript presents a reproducible domain-aware reliability-audit framework for infrastructure image detection and demonstrates it through a validation-boundary study of RDD2022 road-damage detection. The paper addresses validation-target mismatch in civil-infrastructure image AI: ordinary mixed-domain validation can describe recognition under shared dataset context, while domain-withheld validation is needed before detector scores are interpreted as new-region screening evidence.

The numerical evidence combines a seven-domain RDD2022 audit, five-seed YOLOv8n ordinary-versus-LODO validation, a YOLOv8s source-domain budget sweep, detector-family checks with Faster R-CNN, RetinaNet, and RT-DETR-L, class/domain diagnostics, calibration diagnostics, and confidence-frontier analysis. RT-DETR-L improves the subset endpoint to 0.3489 ordinary and 0.2635 mean LODO mAP50, but India remains weak at 0.0668 mAP50. The contribution is therefore not a detector leaderboard; it is a technical framework and six-boundary reporting standard for interpreting detector scores as civil-infrastructure screening evidence.

We believe the manuscript fits the Journal of Computing in Civil Engineering because it treats infrastructure image detection as a computing and validation-evidence problem with direct implications for inspection screening, maintenance prioritization, and accountable AI reporting. The reproducibility package is available at {REPO_URL}; raw RDD2022 archives are not redistributed.

The work is original, is not under consideration elsewhere, and reports no specific funding and no competing interests. The author remains responsible for all claims, code, analyses, sources, and final text. AI-assisted tools were used only for workflow orchestration, code drafting, manuscript drafting, and internal review support. No suggested reviewers are provided.

Sincerely,

Lijian REN

Corresponding author: Lijian REN, renlijian@imut.edu.cn
"""
    md.write_text(text, encoding="utf-8")
    tmp = ROUND / "cover_v22_tmp.md"
    tmp.write_text(text, encoding="utf-8")
    v19.add_markdown_to_docx(tmp, docx)
    return md, docx


def build_text() -> str:
    text = V21.read_text(encoding="utf-8")
    text = re.sub(r"^# .*$", "# Domain-Aware Reliability Auditing for Infrastructure Image Detection: A Validation-Boundary Study of RDD2022 Road-Damage Detection", text, count=1, flags=re.M)
    text = re.sub(r"## Abstract\n\n.*?\n\n## Keywords", "## Abstract\n\n" + abstract() + "\n\n## Keywords", text, flags=re.S)
    text = re.sub(
        r"## Keywords\n\n.*?\n\n## 1\. Introduction",
        "## Keywords\n\nCivil infrastructure computing; Road damage detection; Domain shift; Reliability audit; Object detection; Validation boundary; Infrastructure image analysis\n\n## 1. Introduction",
        text,
        flags=re.S,
    )
    text = replace_between(text, "## 1. Introduction", "## 2. Data and task boundary", introduction())
    text = text.replace("## 2. Data and task boundary", "## 2. Dataset, Domain, and Label-Boundary Audit")
    text = re.sub(r"\n!\[Figure 1\..*?fig02_inventory_label_boundary_enhanced\.png\)\n", "\n", text)
    text = replace_between(text, "## 3. Evaluation design", "## 4. Results", framework_section())
    text = text.replace("## 4. Results", "## 4. Reliability-Boundary Results")
    heading_replacements = {
        "### 4.1 Five-seed cross-domain gap": "### 4.1 Mixed-domain validation overstates domain-withheld evidence",
        "### 4.2 Training budget improves scores but does not equalize domains": "### 4.2 Source-domain scaling improves scores but does not remove the validation boundary",
        "### 4.3 Domain response remains heterogeneous": "### 4.3 Held-out domains have non-interchangeable response profiles",
        "### 4.4 Detector-family checks strengthen the validation-boundary result": "### 4.4 Detector-family checks preserve the boundary pattern",
        "### 4.5 Domain diagnostics explain what the aggregate score hides": "### 4.5 Domain diagnostics explain what the aggregate score hides",
        "### 4.6 Canonical per-class error audit identifies D40 as the most fragile observed label": "### 4.6 Class-level prediction exports localize weak evidence",
        "### 4.7 Calibration and domain-descriptor screening sharpen the decision boundary": "### 4.7 Calibration diagnostics qualify detector confidence",
        "### 4.8 Confidence thresholds do not become a referral policy": "### 4.8 Confidence frontiers limit screening interpretation",
        "### 4.9 Evidence boundary for infrastructure screening interpretation": "### 4.9 Supported and unsupported engineering claims",
    }
    for old, new in heading_replacements.items():
        text = text.replace(old, new)
    text = text.replace(
        "The five-seed YOLOv8n baseline shows that ordinary validation is more favorable than LODO validation under the same subset policy.",
        "The first reliability question is whether mixed-domain validation and domain-withheld validation produce the same engineering evidence. They do not. The five-seed YOLOv8n baseline shows that ordinary validation is more favorable than LODO validation under the same subset policy.",
    )
    text = text.replace(
        "The YOLOv8s budget sweep tests whether the domain-sensitivity pattern disappears when more source-domain images are sampled.",
        "The second reliability question is whether source-domain scaling removes the validation boundary. The YOLOv8s budget sweep tests whether the domain-sensitivity pattern disappears when more source-domain images are sampled.",
    )
    text = text.replace(
        "The R33/R34 detector-family checks replace the earlier one-epoch Faster R-CNN probe with a completed eight-epoch Faster R-CNN run, add a four-epoch RetinaNet minimal validation, and add an eight-epoch RT-DETR-L transformer-family validation.",
        "The detector-family question is whether the validation-boundary pattern is an artifact of one detector family. The R33/R34 checks replace the earlier one-epoch Faster R-CNN probe with a completed eight-epoch Faster R-CNN run, add a four-epoch RetinaNet minimal validation, and add an eight-epoch RT-DETR-L transformer-family validation.",
    )
    text = replace_between(text, "## 5. Discussion", "## 6. Limitations", discussion())
    text = replace_between(text, "## 6. Limitations", "## 7. Conclusion", limitations())
    text = replace_between(text, "## 7. Conclusion", "## Declarations", conclusion_and_practical())
    return text


def official_check() -> Path:
    path = ROUND / "JCICE_OFFICIAL_SCOPE_CHECK.md"
    text = f"""# JCICE Official Scope Check

Date checked: {TODAY}

## Sources

| source | relevance |
| --- | --- |
| ASCE Computing Division publications page: https://www.asce.org/communities/institutes-and-technical-groups/computing-division/publications | Official ASCE page listing Journal of Computing in Civil Engineering under Computing Division journals. |
| ASCE Computing in Civil Engineering 2025 proceedings page: https://sp360.asce.org/PersonifyEbusiness/Merchandise/Product-Details/productId/327844678 | Official ASCE page describing computing-in-civil-engineering topics including artificial intelligence in construction and infrastructure, reality capture, data/sensing, and computational technologies. |

## Editorial implication for v22

- Reframe the paper as a Technical Paper on civil-infrastructure computing and validation evidence.
- Make road-damage detection the reproducible case, not the main novelty claim.
- Add a framework figure, practical applications section, and explicit limitations around non-leaderboard detector-family checks.
- Before live submission, verify the Editorial Manager fields and any ASCE author-guide requirements that are not visible from the public scope pages.
"""
    path.write_text(text, encoding="utf-8")
    return path


def readme(outputs: dict[str, Path]) -> Path:
    path = PKG / "00-README-before-JCICE-submission.md"
    text = f"""# 00-README-before-JCICE-submission

Generated: {TODAY}

## Target

Journal of Computing in Civil Engineering, Technical Paper route.

## Current Files

| item | file |
| --- | --- |
| Manuscript Markdown | `manuscript_v22_jcice_technical_paper.md` |
| Manuscript DOCX | `Manuscript_JCICE_v22_technical_paper.docx` |
| Cover letter DOCX | `CoverLetter_JCICE_v22_technical_paper.docx` |
| Framework Figure 1 | `figures_enhanced/fig01_jcice_domain_aware_reliability_audit_framework.png` |
| Official scope check | `rounds/R35_jcice_technical_reframe_{TODAY}/JCICE_OFFICIAL_SCOPE_CHECK.md` |

## Compared With AEI v21

- Changed the title and abstract to a Technical Paper framing.
- Rewrote the Introduction around validation-target mismatch as a civil-infrastructure computing problem.
- Rebuilt Section 3 as a domain-aware reliability-audit framework.
- Replaced Figure 1 with a framework figure rather than a dataset-statistics figure.
- Reframed Results by reliability questions.
- Rewrote Discussion with JCICE-oriented subsections.
- Added a Practical Applications section.
- Preserved all audited numeric results from v21; no new unverified numbers were introduced.

## Human Check Before Submission

- Verify live ASCE Editorial Manager fields and any current ASCE author-guide requirements.
- Check generated PDF table wrapping.
- Confirm that the public GitHub package remains free of raw RDD2022 files, trained weights, manuscripts, cover letters, and internal rounds/logs.
"""
    path.write_text(text, encoding="utf-8")
    return path


def validate(text: str) -> dict[str, object]:
    scan_hits = v19.scan_text(text)
    scan_path = PKG / "TEXT_SCAN_v22.md"
    if scan_hits:
        scan_path.write_text("# V22 Text Scan\n\nStatus: FAIL\n\n" + "\n".join(f"- {x}" for x in scan_hits) + "\n", encoding="utf-8")
    else:
        scan_path.write_text("# V22 Text Scan\n\nStatus: PASS; no banned phrases, stale placeholders, or internal strategy language found.\n", encoding="utf-8")
    cite = v19.citation_check(text)
    cite_path = ROUND / "V22_CITATION_CHECK.md"
    cite_path.write_text(
        "# V22 Citation Check\n\n"
        + v19.md_table(
            [
                {"check": "missing reference for citation", "value": ", ".join(map(str, cite["missing"])) or "none"},
                {"check": "uncited reference", "value": ", ".join(map(str, cite["uncited"])) or "none"},
                {"check": "references", "value": cite["refs"]},
                {"check": "cited references", "value": cite["cited"]},
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    fig_df = v19.figure_check(text)
    fig_path = ROUND / "V22_FIGURE_FILE_CHECK.md"
    fig_path.write_text("# V22 Figure File Check\n\n" + v19.md_table(fig_df.to_dict("records")) + "\n", encoding="utf-8")
    abstract_match = re.search(r"## Abstract\s+(.*?)\n\n## Keywords", text, flags=re.S)
    from docx import Document

    doc = Document(V22_DOCX)
    return {
        "scan_path": scan_path,
        "cite_path": cite_path,
        "fig_check_path": fig_path,
        "scan_hits": len(scan_hits),
        "missing_refs": len(cite["missing"]),
        "uncited_refs": len(cite["uncited"]),
        "references": cite["refs"],
        "cited_refs": cite["cited"],
        "figures": len(fig_df),
        "figure_missing": int((~fig_df["exists"]).sum()) if not fig_df.empty else 0,
        "abstract_words": len(re.findall(r"\b[\w.-]+\b", abstract_match.group(1))) if abstract_match else 0,
        "docx_paragraphs": len(doc.paragraphs),
        "docx_tables": len(doc.tables),
        "docx_figures": len(doc.inline_shapes),
    }


def audit(outputs: dict[str, Path], validation: dict[str, object]) -> Path:
    path = ROUND / "R35_AUDIT_REPORT.md"
    text = f"""# R35 Audit Report: JCICE Technical Paper Reframe

Date: {TODAY}

## Target Reframe

| item | status |
| --- | --- |
| Title changed to reliability auditing / validation-boundary framing | Done |
| Abstract changed from detector-result summary to Technical Paper framing | Done |
| Introduction rewritten around validation-target mismatch | Done |
| Section 3 rebuilt as domain-aware reliability-audit framework | Done |
| Figure 1 rebuilt as framework figure | Done |
| Results headings reframed by reliability questions | Done |
| Discussion rewritten for civil-infrastructure computing implications | Done |
| Limitations strengthened for non-leaderboard, non-adaptation, single-dataset, low-score boundaries | Done |
| Practical Applications section added | Done |

## Validation

| check | value |
| --- | ---: |
| Text scan hits | {validation["scan_hits"]} |
| Missing references | {validation["missing_refs"]} |
| Uncited references | {validation["uncited_refs"]} |
| Figure references | {validation["figures"]} |
| Missing figure files | {validation["figure_missing"]} |
| Abstract words | {validation["abstract_words"]} |
| DOCX paragraphs | {validation["docx_paragraphs"]} |
| DOCX tables | {validation["docx_tables"]} |
| DOCX figures | {validation["docx_figures"]} |

## Numeric Boundary

No new experimental numbers were introduced. The v22 manuscript reuses v21 audited numbers, including RT-DETR-L ordinary mAP50 0.3489, mean LODO mAP50 0.2635, and India held-out mAP50 0.0668.

## Outputs

{v19.md_table([{"output": k, "path": str(v)} for k, v in outputs.items()])}
"""
    path.write_text(text, encoding="utf-8")
    return path


def state(outputs: dict[str, Path], validation: dict[str, object]) -> Path:
    path = ROUND / "round_state.md"
    text = f"""# R35 JCICE Technical Paper Reframe State

Date: {TODAY}

## Completed Increment Ledger

| requested increment | status | evidence |
| --- | --- | --- |
| Reframe for Journal of Computing in Civil Engineering Technical Paper | Done | v22 title, abstract, introduction, Section 3, discussion, and practical applications |
| Rebuild Figure 1 as framework figure | Done | `fig01_jcice_domain_aware_reliability_audit_framework.png/svg` |
| Preserve audited numeric evidence | Done | v22 uses v21 numbers and R34 audit sources |
| Produce target-specific package | Done | `submission_package/JCICE_RoadDamageDG_2026-05-22` |

## Deferred / Not Done

| item | status | reason |
| --- | --- | --- |
| Domain-adaptation baseline | Deferred | User plan explicitly said not to add more experiments now; safer to preserve reliability-audit scope |
| Multi-seed RT-DETR-L | Deferred | Larger compute round; not necessary for JCICE narrative reframe |
| Live Editorial Manager field-by-field compliance | Partial | Public scope checked; exact EM fields must be verified at upload |

## Validation

| check | value |
| --- | ---: |
| Text scan hits | {validation["scan_hits"]} |
| Missing references | {validation["missing_refs"]} |
| Missing figure files | {validation["figure_missing"]} |
| Abstract words | {validation["abstract_words"]} |
"""
    path.write_text(text, encoding="utf-8")
    return path


def main() -> None:
    prepare_package()
    fig1 = make_framework_figure()
    official = official_check()
    text = build_text()
    V22.write_text(text, encoding="utf-8")
    v19.add_markdown_to_docx(V22, V22_DOCX)
    cover_md, cover_docx = cover_letter()
    outputs = {
        "v22 manuscript": V22,
        "v22 DOCX": V22_DOCX,
        "JCICE cover letter markdown": cover_md,
        "JCICE cover letter DOCX": cover_docx,
        "framework figure PNG": fig1,
        "framework figure SVG": fig1.with_suffix(".svg"),
        "official scope check": official,
    }
    readme_path = readme(outputs)
    outputs["JCICE README"] = readme_path
    validation = validate(text)
    outputs["text scan"] = validation["scan_path"]
    outputs["citation check"] = validation["cite_path"]
    outputs["figure file check"] = validation["fig_check_path"]
    audit_path = audit(outputs, validation)
    outputs["R35 audit report"] = audit_path
    state_path = state(outputs, validation)
    outputs["round state"] = state_path
    for key, value in outputs.items():
        print(f"{key}={value}")
    for key, value in validation.items():
        if isinstance(value, Path):
            continue
        print(f"{key}={value}")


if __name__ == "__main__":
    main()
